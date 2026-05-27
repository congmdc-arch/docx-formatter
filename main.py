"""
DocX Formatter Backend
Chuẩn hóa định dạng đề án tốt nghiệp theo template cố định.

Run: uvicorn main:app --reload --port 8000
"""

import io
import re
import copy
from pathlib import Path
from typing import Optional

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

app = FastAPI(title="DocX Formatter")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ============================================================
# TEMPLATE STYLE CONFIG
# ============================================================
STYLE_CONFIG = {
    "page": {
        "width_cm": 21.0,
        "height_cm": 29.7,
        "margin_top_cm": 3.5,
        "margin_bottom_cm": 3.0,
        "margin_left_cm": 3.5,
        "margin_right_cm": 2.0,
    },
    "default": {
        "font": "Times New Roman",
        "size_pt": 14,
        "line_spacing": 1.5,
        "space_before_pt": 0,
        "space_after_pt": 6,
        "align": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "first_line_indent_cm": 1.27,
    },
    "heading1": {
        # Chương X + tên chương (2 dòng liên tiếp)
        "font": "Times New Roman",
        "size_pt": 14,
        "bold": True,
        "align": WD_ALIGN_PARAGRAPH.CENTER,
        "caps": True,
        "space_before_pt": 0,
        "space_after_pt": 0,
        "line_spacing": 1.5,
        "first_line_indent_cm": None,
    },
    "heading2": {
        # 1.1. TÊN MỤC
        "font": "Times New Roman",
        "size_pt": 14,
        "bold": True,
        "align": WD_ALIGN_PARAGRAPH.LEFT,
        "caps": True,
        "space_before_pt": 0,
        "space_after_pt": 0,
        "line_spacing": 1.5,
        "first_line_indent_cm": None,
    },
    "heading3": {
        # 1.1.1. Tên tiểu mục
        "font": "Times New Roman",
        "size_pt": 14,
        "bold": True,
        "italic": False,
        "align": WD_ALIGN_PARAGRAPH.LEFT,
        "space_before_pt": 0,
        "space_after_pt": 0,
        "line_spacing": 1.5,
    },
    "heading4": {
        "font": "Times New Roman",
        "size_pt": 14,
        "bold": True,
        "italic": True,
        "align": WD_ALIGN_PARAGRAPH.LEFT,
        "space_before_pt": 0,
        "space_after_pt": 0,
        "line_spacing": 1.5,
        "first_line_indent_cm": None,
    },
    "heading5": {
        "font": "Times New Roman",
        "size_pt": 14,
        "bold": False,
        "italic": True,
        "align": WD_ALIGN_PARAGRAPH.LEFT,
        "space_before_pt": 0,
        "space_after_pt": 0,
        "line_spacing": 1.5,
        "first_line_indent_cm": None,
    },
    "caption_figure": {
        # Hình X.Y. Tên hình — dưới hình, căn giữa
        "font": "Times New Roman",
        "size_pt": 13,
        "bold": False,
        "italic": True,
        "align": WD_ALIGN_PARAGRAPH.CENTER,
        "space_before_pt": 0,
        "space_after_pt": 0,
        "line_spacing": 1.5,
    },
    "caption_table": {
        # Bảng X.Y. Tên bảng — trên bảng, căn giữa
        "font": "Times New Roman",
        "size_pt": 13,
        "bold": True,
        "italic": False,
        "align": WD_ALIGN_PARAGRAPH.CENTER,
        "space_before_pt": 0,
        "space_after_pt": 0,
        "line_spacing": 1.5,
    },
    "table_content": {
        "font": "Times New Roman",
        "size_pt": 13,
    },
    "list": {
        "font": "Times New Roman",
        "size_pt": 14,
        "line_spacing": 1.5,
    },
}

# ============================================================
# HELPERS
# ============================================================

def set_run_font(run, font_name: str, size_pt: float,
                 bold: Optional[bool] = None,
                 italic: Optional[bool] = None,
                 color: Optional[RGBColor] = None,
                 caps: bool = False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    # Force Times New Roman for East Asian (Vietnamese) chars
    rpr = run._r.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)

    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    else:
        # Reset màu về tự động (xóa highlight, màu thừa)
        run.font.color.theme_color = None
        rpr2 = run._r.get_or_add_rPr()
        color_el = rpr2.find(qn('w:color'))
        if color_el is not None:
            color_el.set(qn('w:val'), 'auto')
        # Xóa highlight
        hl = rpr2.find(qn('w:highlight'))
        if hl is not None:
            rpr2.remove(hl)
    if caps:
        run.font.all_caps = True
    else:
        run.font.all_caps = None


def set_paragraph_format(para, cfg: dict):
    pf = para.paragraph_format
    if 'align' in cfg:
        pf.alignment = cfg['align']
    if 'line_spacing' in cfg:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = cfg['line_spacing']
    if 'space_before_pt' in cfg:
        pf.space_before = Pt(cfg['space_before_pt'])
    if 'space_after_pt' in cfg:
        pf.space_after = Pt(cfg['space_after_pt'])
    if 'first_line_indent_cm' in cfg and cfg['first_line_indent_cm'] is not None:
        pf.first_line_indent = Cm(cfg['first_line_indent_cm'])
    else:
        pf.first_line_indent = None


def format_paragraph_runs(para, cfg: dict):
    for run in para.runs:
        set_run_font(
            run,
            font_name=cfg.get('font', 'Times New Roman'),
            size_pt=cfg.get('size_pt', 14),
            bold=cfg.get('bold'),
            italic=cfg.get('italic'),
            caps=cfg.get('caps', False),
        )


# ============================================================
# CAPTION NUMBER TRACKER
# ============================================================

class CaptionTracker:
    """
    Track hình/bảng theo chương.
    Hình X.Y — đánh số tăng dần trong chương X
    Bảng X.Y — tương tự
    """
    def __init__(self):
        self.current_chapter = 0
        self.fig_count = 0
        self.tbl_count = 0

    def set_chapter(self, chapter_num: int):
        if chapter_num != self.current_chapter:
            self.current_chapter = chapter_num
            self.fig_count = 0
            self.tbl_count = 0

    def next_figure(self) -> str:
        self.fig_count += 1
        return f"Hình {self.current_chapter}.{self.fig_count}"

    def next_table(self) -> str:
        self.tbl_count += 1
        return f"Bảng {self.current_chapter}.{self.tbl_count}"


# Regex nhận dạng caption cũ
RE_FIGURE_CAPTION = re.compile(
    r'^(Hình\s+\d+[\.\s]*\d*\.?\s*[:\-]?\s*)', re.IGNORECASE
)
RE_TABLE_CAPTION = re.compile(
    r'^(Bảng\s+\d+[\.\s]*\d*\.?\s*[:\-]?\s*)', re.IGNORECASE
)
# Nhận dạng heading chương
RE_CHAPTER = re.compile(r'^Chương\s+(\d+)', re.IGNORECASE)
RE_HEADING2 = re.compile(r'^\d+\.\d+\.?\s')
RE_HEADING3 = re.compile(r'^\d+\.\d+\.\d+\.?\s')


def detect_chapter_from_text(text: str) -> Optional[int]:
    m = RE_CHAPTER.match(text.strip())
    if m:
        return int(m.group(1))
    return None


def reformat_caption_text(text: str, tracker: CaptionTracker, is_figure: bool) -> str:
    """Chuẩn hóa lại số thứ tự caption"""
    if is_figure:
        prefix = tracker.next_figure()
        clean = RE_FIGURE_CAPTION.sub('', text).strip()
        return f"{prefix}. {clean}" if clean else prefix
    else:
        prefix = tracker.next_table()
        clean = RE_TABLE_CAPTION.sub('', text).strip()
        return f"{prefix}. {clean}" if clean else prefix


# ============================================================
# MAIN FORMAT FUNCTION
# ============================================================

def format_document(doc: Document) -> Document:
    cfg = STYLE_CONFIG
    
    # Tìm index của Heading 1 đầu tiên
    first_heading_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'Heading 1' and p.text.strip():
            first_heading_idx = i
            break

    tracker = CaptionTracker()

    # 1. Page layout
    for section in doc.sections:
        section.page_width = Cm(cfg["page"]["width_cm"])
        section.page_height = Cm(cfg["page"]["height_cm"])
        section.top_margin = Cm(cfg["page"]["margin_top_cm"])
        section.bottom_margin = Cm(cfg["page"]["margin_bottom_cm"])
        section.left_margin = Cm(cfg["page"]["margin_left_cm"])
        section.right_margin = Cm(cfg["page"]["margin_right_cm"])

    # 2. Format paragraphs
    for para in doc.paragraphs:
        style_name = para.style.name
        text = para.text.strip()

        # ---- Trang bìa: trước Heading 1 đầu tiên ----
        if first_heading_idx and i < first_heading_idx:
            pf = para.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = None
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = 1.5
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                rpr = run._r.get_or_add_rPr()
                rFonts = rpr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rpr.insert(0, rFonts)
                for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
                    rFonts.set(qn(attr), 'Times New Roman')
            continue  # ← skip các xử lý bên dưới

        # Detect chapter để cập nhật tracker
        ch = detect_chapter_from_text(text)
        if ch:
            tracker.set_chapter(ch)

        # ---- Heading 1 ----
        if style_name == 'Heading 1':
            h_cfg = cfg['heading1']
            set_paragraph_format(para, h_cfg)
            format_paragraph_runs(para, h_cfg)

        # ---- Heading 2 ----
        elif style_name == 'Heading 2':
            h_cfg = cfg['heading2']
            set_paragraph_format(para, h_cfg)
            format_paragraph_runs(para, h_cfg)

        # ---- Heading 3 ----
        elif style_name == 'Heading 3':
            h_cfg = cfg['heading3']
            set_paragraph_format(para, h_cfg)
            format_paragraph_runs(para, h_cfg)

        # ---- Heading 4 ----
        elif style_name == 'Heading 4':
            h_cfg = cfg['heading4']
            set_paragraph_format(para, h_cfg)
            format_paragraph_runs(para, h_cfg)

        # ---- Heading 5 ----
        elif style_name == 'Heading 5':
            h_cfg = cfg['heading5']
            set_paragraph_format(para, h_cfg)
            format_paragraph_runs(para, h_cfg)

        # ---- Caption (hình/bảng) ----
        elif style_name == 'Caption':
            is_fig = bool(RE_FIGURE_CAPTION.match(text))
            is_tbl = bool(RE_TABLE_CAPTION.match(text))

            if is_fig:
                new_text = reformat_caption_text(text, tracker, is_figure=True)
                h_cfg = cfg['caption_figure']
            elif is_tbl:
                new_text = reformat_caption_text(text, tracker, is_figure=False)
                h_cfg = cfg['caption_table']
            else:
                new_text = text
                h_cfg = cfg['caption_figure']  # default

            # Rewrite text vào run đầu, xóa runs thừa
            if para.runs:
                para.runs[0].text = new_text
                for run in para.runs[1:]:
                    run.text = ''
            set_paragraph_format(para, h_cfg)
            format_paragraph_runs(para, h_cfg)

        # ---- List Paragraph ----
        elif style_name == 'List Paragraph':
            l_cfg = cfg['list']
            set_paragraph_format(para, {
                'align': WD_ALIGN_PARAGRAPH.JUSTIFY,
                'line_spacing': l_cfg['line_spacing'],
                'space_before_pt': 0,
                'space_after_pt': 3,
            })
            format_paragraph_runs(para, l_cfg)

        # ---- Normal / Body Text / default ----
        elif style_name in ('Normal', 'Body Text', 'Default Paragraph Font'):
            if not text:
                continue  # empty paragraph — skip format
            d_cfg = cfg['default']
            # Check nếu là paragraph thường (có indent) hay special (trang bìa, etc.)
            pf = para.paragraph_format
            has_indent = (
                para.alignment in (WD_ALIGN_PARAGRAPH.JUSTIFY, None)
                or para.alignment == WD_ALIGN_PARAGRAPH.LEFT
            )
            set_paragraph_format(para, {
                'align': WD_ALIGN_PARAGRAPH.JUSTIFY,
                'line_spacing': d_cfg['line_spacing'],
                'space_before_pt': d_cfg['space_before_pt'],
                'space_after_pt': d_cfg['space_after_pt'],
                'first_line_indent_cm': d_cfg['first_line_indent_cm'] if has_indent else None,
            })
            format_paragraph_runs(para, {
                'font': d_cfg['font'],
                'size_pt': d_cfg['size_pt'],
                'bold': None,    # giữ bold gốc nếu có
                'italic': None,  # giữ italic gốc nếu có
            })
            # Chỉ reset color/highlight, giữ bold/italic gốc của từng run
            for run in para.runs:
                run.font.name = d_cfg['font']
                run.font.size = Pt(d_cfg['size_pt'])
                # Force font cho tiếng Việt
                rpr = run._r.get_or_add_rPr()
                rFonts = rpr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rpr.insert(0, rFonts)
                for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
                    rFonts.set(qn(attr), d_cfg['font'])
                # Xóa highlight và màu
                hl = rpr.find(qn('w:highlight'))
                if hl is not None:
                    rpr.remove(hl)
                color_el = rpr.find(qn('w:color'))
                if color_el is not None:
                    color_el.set(qn('w:val'), 'auto')

    # 3. Format tables
    for table in doc.tables:
        _format_table(table, cfg)

    # 4. Update default style trong document
    _update_doc_defaults(doc, cfg)

    return doc


def _format_table(table, cfg: dict):
    tc_cfg = cfg['table_content']
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                pf = para.paragraph_format
                pf.space_before = Pt(2)
                pf.space_after = Pt(2)
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                pf.line_spacing = 1.2
                for run in para.runs:
                    run.font.name = tc_cfg['font']
                    run.font.size = Pt(tc_cfg['size_pt'])
                    rpr = run._r.get_or_add_rPr()
                    rFonts = rpr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = OxmlElement('w:rFonts')
                        rpr.insert(0, rFonts)
                    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
                        rFonts.set(qn(attr), tc_cfg['font'])
                    # Xóa highlight
                    hl = rpr.find(qn('w:highlight'))
                    if hl is not None:
                        rpr.remove(hl)


def _update_doc_defaults(doc: Document, cfg: dict):
    """Cập nhật docDefaults để font mặc định áp dụng toàn document"""
    styles_part = doc.part.styles
    styles_xml = styles_part._element
    doc_defaults = styles_xml.find(qn('w:docDefaults'))
    if doc_defaults is None:
        return

    rpr_default = doc_defaults.find('.//' + qn('w:rPrDefault'))
    if rpr_default is None:
        return

    rpr = rpr_default.find(qn('w:rPr'))
    if rpr is None:
        rpr = OxmlElement('w:rPr')
        rpr_default.append(rpr)

    # Set font
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.insert(0, rFonts)
    font = cfg['default']['font']
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), font)

    # Set size (half-points)
    sz = rpr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rpr.append(sz)
    sz.set(qn('w:val'), str(int(cfg['default']['size_pt'] * 2)))

    szCs = rpr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rpr.append(szCs)
    szCs.set(qn('w:val'), str(int(cfg['default']['size_pt'] * 2)))


# ============================================================
# API ENDPOINTS
# ============================================================

@app.get("/")
def root():
    return {"status": "ok", "service": "DocX Formatter"}


@app.post("/format")
async def format_docx(file: UploadFile = File(...)):
    if not file.filename.endswith('.docx'):
        raise HTTPException(400, "Chỉ hỗ trợ file .docx")

    content = await file.read()
    try:
        doc = Document(io.BytesIO(content))
    except Exception as e:
        raise HTTPException(400, f"Không đọc được file: {e}")

    import traceback
    try:
        doc = format_document(doc)
    except Exception as e:
        raise HTTPException(500, f"Lỗi format: {traceback.format_exc()}")

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    from urllib.parse import quote
    filename = file.filename.replace('.docx', '_formatted.docx')
    encoded_filename = quote(filename)
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )


@app.post("/inspect")
async def inspect_docx(file: UploadFile = File(...)):
    """Endpoint để debug — trả về cấu trúc file"""
    if not file.filename.endswith('.docx'):
        raise HTTPException(400, "Chỉ hỗ trợ file .docx")

    content = await file.read()
    doc = Document(io.BytesIO(content))

    paragraphs = []
    for i, p in enumerate(doc.paragraphs[:100]):
        if p.text.strip():
            paragraphs.append({
                "index": i,
                "style": p.style.name,
                "text": p.text[:120],
                "align": str(p.alignment),
                "runs_count": len(p.runs),
            })

    section = doc.sections[0]
    return {
        "filename": file.filename,
        "page": {
            "width_cm": round(section.page_width.cm, 2),
            "height_cm": round(section.page_height.cm, 2),
            "margin_top": round(section.top_margin.cm, 2),
            "margin_bottom": round(section.bottom_margin.cm, 2),
            "margin_left": round(section.left_margin.cm, 2),
            "margin_right": round(section.right_margin.cm, 2),
        },
        "tables": len(doc.tables),
        "paragraphs_total": len(doc.paragraphs),
        "paragraphs_sample": paragraphs,
    }
